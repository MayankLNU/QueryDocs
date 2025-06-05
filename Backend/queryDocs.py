import os
from openai import AzureOpenAI
from dotenv import load_dotenv
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import AzureOpenAIEmbeddings
from langchain_community.vectorstores.azuresearch import AzureSearch
from langchain_core.documents import Document


load_dotenv()


AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_CHAT_DEPLOYMENT_NAME = os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME")
AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME")
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION")
AZURE_SEARCH_ENDPOINT = os.getenv("AZURE_SEARCH_ENDPOINT")
AZURE_SEARCH_ADMIN_KEY = os.getenv("AZURE_SEARCH_ADMIN_KEY")
AZURE_SEARCH_INDEX_NAME = os.getenv("AZURE_SEARCH_INDEX_NAME")
 

if not all([AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_CHAT_DEPLOYMENT_NAME,
            AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME, AZURE_OPENAI_API_VERSION, AZURE_SEARCH_ENDPOINT, 
            AZURE_SEARCH_ADMIN_KEY, AZURE_SEARCH_INDEX_NAME]):
    raise ValueError("Azure OpenAI environment variables not set. Please check your .env file.")
 

chat_client = AzureOpenAI(
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
    azure_endpoint=AZURE_OPENAI_ENDPOINT
)
 
embeddings_model = AzureOpenAIEmbeddings(
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    deployment=AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME,
    openai_api_key=AZURE_OPENAI_API_KEY,
    openai_api_version=AZURE_OPENAI_API_VERSION,
    chunk_size=1024
)


vector_store = AzureSearch(
    azure_search_endpoint=AZURE_SEARCH_ENDPOINT,
    azure_search_key=AZURE_SEARCH_ADMIN_KEY,
    index_name=AZURE_SEARCH_INDEX_NAME,
    embedding_function=embeddings_model.embed_query
)


def extract_text_and_file_extension(file, file_extension):
    
    if file_extension == '.pdf':
        reader = PyPDF2.PdfReader(file)
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    
    if file_extension == '.docx':
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs]).strip()


def process_document_for_rag(filename, file_extension, extracted_text):
 
    # Chunk the extracted text
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
        separators=["\n\n", "\n", " ", ""]
    )
    chunks = text_splitter.split_text(extracted_text)
 
    # Prepare documents for vector store with metadata
    docs_to_index = [
        Document(
            page_content=chunk,
            metadata={
                "source": filename,
                "chunk_id": i,
                "file_type": file_extension
            }
        ) for i, chunk in enumerate(chunks)
    ]
 
    # Add documents to Azure AI Search (generates embeddings internally)
    vector_store.add_documents(docs_to_index)
    return True


def answer_question_from_docs(question):
    retrieved_docs = vector_store.similarity_search(query=question, k=3)
 
    if not retrieved_docs:
        return "No relevant information found in the indexed documents for your question."
 
    context_for_llm = "\n\n".join([doc.page_content for doc in retrieved_docs])
 
    messages = [
        {"role": "system", "content": "You are a helpful assistant. Answer the question based ONLY on the provided context. If the answer is not in the context, say you don't know."},
        {"role": "user", "content": f"Context:\n{context_for_llm}\n\nQuestion: {question}"}
    ]
 
    response = chat_client.chat.completions.create(
        model=AZURE_OPENAI_CHAT_DEPLOYMENT_NAME,
        messages=messages,
        temperature=0.6,
        max_tokens=500
    )
 
    return response.choices[0].message.content.strip()